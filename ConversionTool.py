import bibtexparser
from bibtexparser.bparser import BibTexParser
from docx import Document
import os
import tkinter as tk
from tkinter import filedialog, messagebox, Checkbutton, BooleanVar

def categorize_bibtex_entries(bibtex_file, output_dir, include_publications, include_conferences, include_patents):
    # Read the .bibtex file
    with open(bibtex_file, 'r', encoding='utf-8') as f:
        parser = BibTexParser(common_strings=True)
        bib_database = bibtexparser.load(f, parser=parser)

    # Initialize lists for publications, conferences, and patents
    publications = []
    conferences = []
    patents = []

    # Define entry types for publications and conferences
    publication_types = {'article', 'book', 'booklet', 'manual', 'mastersthesis', 'phdthesis', 'techreport'}
    conference_types = {'inproceedings', 'conference'}

    # Categorize entries
    for entry in bib_database.entries:
        # Check if the entry is a patent (using ENTRYTYPE or type field)
        if entry.get('ENTRYTYPE', '').lower() == 'patent' or entry.get('type', '').lower() == 'patent':
            patents.append(entry)
        # Check if the entry is a conference based on ENTRYTYPE or keywords in booktitle/journal
        elif (
            entry['ENTRYTYPE'] in conference_types
            or 'conference' in entry.get('booktitle', '').lower()
            or 'meeting' in entry.get('booktitle', '').lower()
            or 'workshop' in entry.get('booktitle', '').lower()
            or 'symposium' in entry.get('booktitle', '').lower()
            or 'AGU' in entry.get('booktitle', '').upper()
            or 'conference' in entry.get('journal', '').lower()
            or 'meeting' in entry.get('journal', '').lower()
            or 'workshop' in entry.get('journal', '').lower()
            or 'symposium' in entry.get('journal', '').lower()
            or 'AGU' in entry.get('journal', '').upper()
        ):
            conferences.append(entry)
        # Otherwise, treat as a publication
        elif entry['ENTRYTYPE'] in publication_types:
            publications.append(entry)

    # Write publications to a .doc file in APA format
    if include_publications.get() and publications:
        pub_doc_path = os.path.join(output_dir, 'publications.docx')
        pub_doc = Document()
        pub_doc.add_heading('Publications', level=1)
        for entry in publications:
            add_apa_paragraph(pub_doc, entry, is_conference=False)
        pub_doc.save(pub_doc_path)

    # Write conferences to a .doc file in APA format
    if include_conferences.get() and conferences:
        conf_doc_path = os.path.join(output_dir, 'conferences.docx')
        conf_doc = Document()
        conf_doc.add_heading('Conferences', level=1)
        for entry in conferences:
            add_apa_paragraph(conf_doc, entry, is_conference=True)
        conf_doc.save(conf_doc_path)

    # Write patents to a .doc file in APA format
    if include_patents.get() and patents:
        patent_doc_path = os.path.join(output_dir, 'patents.docx')
        patent_doc = Document()
        patent_doc.add_heading('Patents', level=1)
        for entry in patents:
            add_apa_paragraph(patent_doc, entry, is_patent=True)
        patent_doc.save(patent_doc_path)

    messagebox.showinfo("Success", f"Processing complete. Files saved to {output_dir}.")

def add_apa_paragraph(doc, entry, is_conference=False, is_patent=False):
    """Add an APA-formatted paragraph to a Word document with automatic italics."""
    paragraph = doc.add_paragraph()

    # Authors
    if 'author' in entry:
        authors = entry['author'].split(' and ')
        formatted_authors = []
        for i, author in enumerate(authors):
            if ',' in author:
                last, first = author.split(',', 1)
                formatted_authors.append(f"{first.strip()} {last.strip()}")
            else:
                formatted_authors.append(author.strip())
            if i == len(authors) - 2:  # Add "&" before the last author
                formatted_authors.append("&")
        paragraph.add_run(', '.join(formatted_authors) + '. ')

    # Year (and month, if available)
    if 'year' in entry:
        if 'month' in entry:
            paragraph.add_run(f"({entry['year']}, {entry['month']}). ")
        else:
            paragraph.add_run(f"({entry['year']}). ")

    # Title
    if 'title' in entry:
        paragraph.add_run(f"{entry['title']}. ")

    # Journal, Conference, or Patent
    if is_conference:
        # Use 'booktitle' if available, otherwise fall back to 'journal'
        conference_name = entry.get('booktitle', entry.get('journal', ''))
        if conference_name:
            paragraph.add_run("In ")
            run = paragraph.add_run(conference_name)
            run.italic = True  # Italicize conference name
        if 'number' in entry:
            paragraph.add_run(f" (No. {entry['number']}).")
        elif 'pages' in entry:
            paragraph.add_run(f" (pp. {entry['pages']}).")
    elif is_patent:
        if 'journal' in entry:
            paragraph.add_run(f"{entry['journal']}.")
    else:
        if 'journal' in entry:
            run = paragraph.add_run(entry['journal'])
            run.italic = True  # Italicize journal name
        if 'volume' in entry:
            paragraph.add_run(f", {entry['volume']}")
        if 'number' in entry:
            paragraph.add_run(f"({entry['number']})")
        if 'pages' in entry:
            paragraph.add_run(f", {entry['pages']}.")

    # Publisher (for books and conferences)
    if 'publisher' in entry and not is_conference and not is_patent:
        paragraph.add_run(f" {entry['publisher']}.")

    # URL or DOI
    if 'url' in entry:
        paragraph.add_run(f" Retrieved from {entry['url']}")
    elif 'doi' in entry:
        paragraph.add_run(f" https://doi.org/{entry['doi']}")

def select_file():
    file_path = filedialog.askopenfilename(filetypes=[("BibTeX files", "*.bib")])
    if file_path:
        file_entry.delete(0, tk.END)
        file_entry.insert(0, file_path)

def select_output_dir():
    dir_path = filedialog.askdirectory()
    if dir_path:
        output_entry.delete(0, tk.END)
        output_entry.insert(0, dir_path)

def process_file():
    bibtex_file = file_entry.get()
    output_dir = output_entry.get()

    if not bibtex_file or not output_dir:
        messagebox.showerror("Error", "Please select a .bib file and an output directory.")
        return

    if not os.path.exists(bibtex_file):
        messagebox.showerror("Error", "The selected .bib file does not exist.")
        return

    if not os.path.exists(output_dir):
        messagebox.showerror("Error", "The selected output directory does not exist.")
        return

    categorize_bibtex_entries(bibtex_file, output_dir, include_publications, include_conferences, include_patents)

# Create the main window
root = tk.Tk()
root.title("BibTeX to APA Formatter")

# File selection
file_label = tk.Label(root, text="Select .bib file:")
file_label.grid(row=0, column=0, padx=5, pady=5)
file_entry = tk.Entry(root, width=50)
file_entry.grid(row=0, column=1, padx=5, pady=5)
file_button = tk.Button(root, text="Browse", command=select_file)
file_button.grid(row=0, column=2, padx=5, pady=5)

# Output directory selection
output_label = tk.Label(root, text="Select output directory:")
output_label.grid(row=1, column=0, padx=5, pady=5)
output_entry = tk.Entry(root, width=50)
output_entry.grid(row=1, column=1, padx=5, pady=5)
output_button = tk.Button(root, text="Browse", command=select_output_dir)
output_button.grid(row=1, column=2, padx=5, pady=5)

# Toggles for sections
include_publications = BooleanVar(value=True)
include_conferences = BooleanVar(value=True)
include_patents = BooleanVar(value=True)

publications_check = Checkbutton(root, text="Include Publications", variable=include_publications)
publications_check.grid(row=2, column=0, padx=5, pady=5)

conferences_check = Checkbutton(root, text="Include Conferences", variable=include_conferences)
conferences_check.grid(row=2, column=1, padx=5, pady=5)

patents_check = Checkbutton(root, text="Include Patents", variable=include_patents)
patents_check.grid(row=2, column=2, padx=5, pady=5)

# Process button
process_button = tk.Button(root, text="Process", command=process_file)
process_button.grid(row=3, column=1, padx=5, pady=10)

# Run the application
root.mainloop()
